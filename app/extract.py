"""Plain-text extraction from uploaded documents.

A generic capability: turn an uploaded file's bytes into text. Brain stays
content-agnostic — it has no idea whether the document is a resume, a contract,
or a recipe.

PDF and DOCX are first-class. PDFs are extracted with ``pdfplumber`` when
available — it preserves reading order far better than a naive text dump and can
recover **tables**, which are rendered as Markdown so their structure survives
into the chunker and the embeddings. A page's table regions are excluded from the
prose pass so cells aren't duplicated. ``pypdf`` is the fallback when pdfplumber
is unavailable or yields nothing (e.g. an oddly-encoded PDF). Legacy ``.doc`` is
not supported.
"""
from __future__ import annotations

import io
import re

_WS_RE = re.compile(r"[ \t]+")
_BLANKS_RE = re.compile(r"\n{3,}")


class ExtractionError(Exception):
    """Raised when a document cannot be parsed into text."""


def _clean(text: str) -> str:
    text = text.replace("\r\n", "\n").replace("\r", "\n")
    text = _WS_RE.sub(" ", text)
    text = _BLANKS_RE.sub("\n\n", text)
    return text.strip()


def _table_to_markdown(rows) -> str:
    """Render a grid of cells (list[list[str|None]]) as a Markdown pipe table.

    Newlines inside a cell are flattened; fully-empty rows are dropped; short
    rows are padded so every row has the same column count. Returns ``""`` when
    there is nothing worth rendering."""
    if not rows:
        return ""
    cleaned: list[list[str]] = []
    for row in rows:
        cells = [(str(c) if c is not None else "").replace("\n", " ").strip() for c in row]
        if any(cells):
            cleaned.append(cells)
    if not cleaned:
        return ""
    width = max(len(r) for r in cleaned)
    norm = [r + [""] * (width - len(r)) for r in cleaned]
    header = norm[0]
    lines = ["| " + " | ".join(header) + " |", "| " + " | ".join(["---"] * width) + " |"]
    for row in norm[1:]:
        lines.append("| " + " | ".join(row) + " |")
    return "\n".join(lines)


def _extract_pdf_pdfplumber(data: bytes) -> str:
    """Layout-aware PDF text + Markdown tables. Returns "" to signal a fallback."""
    try:
        import pdfplumber
    except Exception:  # noqa: BLE001 — not installed; caller falls back to pypdf
        return ""

    pages_out: list[str] = []
    try:
        with pdfplumber.open(io.BytesIO(data)) as pdf:
            for page in pdf.pages:
                try:
                    tables = page.find_tables() or []
                except Exception:  # noqa: BLE001
                    tables = []
                bboxes = []
                for t in tables:
                    try:
                        bboxes.append(t.bbox)
                    except Exception:  # noqa: BLE001
                        continue

                # Prose pass: exclude words that fall inside a detected table so
                # cell text isn't emitted twice (once as prose, once as a table).
                def _outside_tables(obj, _bboxes=bboxes):
                    cx = (obj.get("x0", 0) + obj.get("x1", 0)) / 2
                    cy = (obj.get("top", 0) + obj.get("bottom", 0)) / 2
                    for x0, top, x1, bottom in _bboxes:
                        if x0 <= cx <= x1 and top <= cy <= bottom:
                            return False
                    return True

                try:
                    source = page.filter(_outside_tables) if bboxes else page
                    page_text = source.extract_text(x_tolerance=1.5) or ""
                except Exception:  # noqa: BLE001
                    try:
                        page_text = page.extract_text() or ""
                    except Exception:  # noqa: BLE001
                        page_text = ""

                parts: list[str] = []
                if page_text.strip():
                    parts.append(page_text)
                for t in tables:
                    try:
                        md = _table_to_markdown(t.extract())
                    except Exception:  # noqa: BLE001
                        md = ""
                    if md:
                        parts.append(md)
                if parts:
                    pages_out.append("\n\n".join(parts))
    except Exception:  # noqa: BLE001 — corrupt/unsupported PDF; let caller fall back
        return ""
    return "\n\n".join(pages_out)


def _extract_pdf_pypdf(data: bytes) -> str:
    try:
        from pypdf import PdfReader
    except Exception as exc:  # noqa: BLE001
        raise ExtractionError("no PDF backend available (install pdfplumber or pypdf)") from exc
    try:
        reader = PdfReader(io.BytesIO(data))
    except Exception as exc:  # noqa: BLE001
        raise ExtractionError(f"could not read PDF: {exc}") from exc
    pages = []
    for page in reader.pages:
        try:
            pages.append(page.extract_text() or "")
        except Exception:  # noqa: BLE001 — skip an unreadable page, keep the rest
            continue
    return "\n\n".join(p for p in pages if p.strip())


def _extract_pdf(data: bytes) -> str:
    # Prefer pdfplumber (better order + tables); fall back to pypdf when it's
    # missing or returns nothing usable.
    text = _extract_pdf_pdfplumber(data)
    if text and text.strip():
        return text
    return _extract_pdf_pypdf(data)


def _extract_docx(data: bytes) -> str:
    try:
        import docx  # python-docx
    except Exception as exc:  # noqa: BLE001
        raise ExtractionError("python-docx is not installed") from exc
    try:
        document = docx.Document(io.BytesIO(data))
    except Exception as exc:  # noqa: BLE001
        raise ExtractionError(f"could not read DOCX: {exc}") from exc
    parts = [p.text for p in document.paragraphs if p.text and p.text.strip()]
    # Render tables as Markdown so structure (rows/columns) survives — docs often
    # lay out key/value data or comparisons in tables.
    for table in document.tables:
        rows = [[c.text.strip() for c in row.cells] for row in table.rows]
        md = _table_to_markdown(rows)
        if md:
            parts.append(md)
    return "\n".join(parts)


def extract_text(data: bytes, *, filename: str = "", content_type: str = "") -> str:
    """Extract plain text from ``data``, dispatching on extension / MIME type."""
    name = (filename or "").lower()
    ctype = (content_type or "").lower()

    if name.endswith(".pdf") or "pdf" in ctype:
        text = _extract_pdf(data)
    elif name.endswith(".docx") or "openxmlformats-officedocument.wordprocessingml" in ctype:
        text = _extract_docx(data)
    elif name.endswith(".doc") or ctype == "application/msword":
        raise ExtractionError("legacy .doc files are not supported — please upload a PDF or DOCX")
    elif name.endswith((".md", ".markdown")) or "markdown" in ctype:
        text = data.decode("utf-8", errors="ignore")
    elif name.endswith(".txt") or ctype.startswith("text/"):
        text = data.decode("utf-8", errors="ignore")
    else:
        # Last-ditch attempt: treat as UTF-8 text.
        try:
            text = data.decode("utf-8")
        except UnicodeDecodeError as exc:
            raise ExtractionError(
                f"unsupported file type (filename='{filename}', content_type='{content_type}')"
            ) from exc

    cleaned = _clean(text)
    if not cleaned:
        raise ExtractionError("no extractable text found in document")
    return cleaned
